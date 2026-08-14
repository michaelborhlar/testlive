"""Turn an uploaded PDF / Word / text document into structured questions.

Everything here runs locally and costs nothing:

* text extraction  — pypdf and python-docx (both open source)
* question parsing — the rule-based reader below, which understands the
  layouts question banks are normally written in
* optional LLM     — if a local Ollama server is running (free, offline),
  it is used to structure messy documents; any failure falls back to the
  rule-based reader, so the feature never depends on a paid service.
"""
import io
import json
import re
import urllib.error
import urllib.request
import uuid

from django.conf import settings
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage

MAX_UPLOAD_BYTES = 25 * 1024 * 1024
SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".txt", ".md")
IMAGE_DIR = "questions"
# Charts compress well, so size in bytes says little — filter on dimensions
# instead, which is what separates a graph from a bullet, rule or logo strip.
MIN_IMAGE_WIDTH = 100
MIN_IMAGE_HEIGHT = 60
MIN_IMAGE_BYTES = 400
# Marker injected into the text stream where a picture appears in the document.
IMAGE_MARKER_RE = re.compile(r"^\[\[IMAGE:(.+?)\]\]$")


def is_meaningful_image(data: bytes) -> bool:
    """True when a picture is large enough to be a figure rather than decoration."""
    if not data or len(data) < MIN_IMAGE_BYTES:
        return False
    try:
        from PIL import Image

        with Image.open(io.BytesIO(data)) as image:
            width, height = image.size
        return width >= MIN_IMAGE_WIDTH and height >= MIN_IMAGE_HEIGHT
    except ImportError:
        return len(data) >= 2048
    except Exception:
        return False


def store_image(data: bytes, extension: str = "png") -> str:
    """Save an extracted picture to media storage, returning its path."""
    extension = (extension or "png").lower().lstrip(".")
    if extension in ("jpeg", "jpe"):
        extension = "jpg"
    name = f"{IMAGE_DIR}/{uuid.uuid4().hex[:12]}.{extension}"
    return default_storage.save(name, ContentFile(data))

# 1.  /  1)  /  Q1.  /  Question 1:
QUESTION_RE = re.compile(r"^\s*(?:q(?:uestion)?\s*)?(\d{1,3})\s*[\.\)\:\-]\s*(.+)$", re.I)
# A.  /  (b)  /  c)
OPTION_RE = re.compile(r"^\s*\(?([a-h])\s*[\.\)\:]\s*(.+)$", re.I)
ANSWER_RE = re.compile(
    r"^\s*(?:correct\s+answer|answer|ans|key)\s*[:\.\-]?\s*(.+)$", re.I
)
EXPLANATION_RE = re.compile(r"^\s*(?:explanation|rationale|reason|why)\s*[:\.\-]?\s*(.+)$", re.I)
MARKS_RE = re.compile(r"[\(\[]\s*(\d{1,2})\s*(?:marks?|points?|pts?)\s*[\)\]]", re.I)
KEY_HEADER_RE = re.compile(r"^\s*(?:answer\s*key|answers|solutions)\s*[:\-]?\s*$", re.I)
KEY_ENTRY_RE = re.compile(r"(\d{1,3})\s*[\.\)\:\-]?\s*([a-h](?:\s*[,&/]\s*[a-h])*)\b", re.I)
CORRECT_MARK_RE = re.compile(r"(?:^\s*\*+\s*|\s*\*+\s*$|\s*\[correct\]\s*|\s*\(correct\)\s*)", re.I)
PAGE_NOISE_RE = re.compile(r"^\s*(?:page\s*)?\d{1,3}\s*(?:of\s*\d{1,3})?\s*$", re.I)


class DocumentError(Exception):
    """Raised when a document cannot be read."""


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------


def extract_text(uploaded_file) -> str:
    name = (uploaded_file.name or "").lower()
    if not name.endswith(SUPPORTED_EXTENSIONS):
        raise DocumentError("Upload a PDF, Word (.docx), text or markdown file.")
    if uploaded_file.size > MAX_UPLOAD_BYTES:
        raise DocumentError("That file is larger than the 10 MB limit.")

    raw = uploaded_file.read()
    if name.endswith(".pdf"):
        return _extract_pdf(raw)
    if name.endswith(".docx"):
        return _extract_docx(raw)
    return raw.decode("utf-8", errors="replace")


def _extract_pdf(raw: bytes) -> str:
    """Page text with an [[IMAGE:path]] marker wherever a picture sits."""
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise DocumentError("PDF support needs 'pypdf' — run pip install pypdf.") from exc

    try:
        reader = PdfReader(io.BytesIO(raw))
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception as exc:
                raise DocumentError("That PDF is password protected.") from exc

        chunks = []
        for page in reader.pages:
            chunks.append(page.extract_text() or "")
            # Graphs and charts are normally printed with the question they
            # belong to, so images are emitted at the end of their own page.
            for marker in _pdf_page_images(page):
                chunks.append(marker)
    except DocumentError:
        raise
    except Exception as exc:
        raise DocumentError(f"The PDF could not be read: {exc}") from exc

    text = "\n".join(chunks).strip()
    if not text:
        raise DocumentError(
            "No text could be extracted — the PDF looks like scanned images. "
            "Export a text-based PDF or paste the questions into a Word file."
        )
    return text


def _pdf_page_images(page):
    markers = []
    try:
        images = list(page.images)
    except Exception:
        return markers

    for image in images:
        try:
            data = image.data
            if not is_meaningful_image(data):
                continue
            extension = (image.name or "png").rsplit(".", 1)[-1]
            markers.append(f"[[IMAGE:{store_image(data, extension)}]]")
        except Exception:
            continue
    return markers


def _extract_docx(raw: bytes) -> str:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover
        raise DocumentError("Word support needs 'python-docx'.") from exc

    try:
        document = docx.Document(io.BytesIO(raw))
    except Exception as exc:
        raise DocumentError(f"The Word document could not be read: {exc}") from exc

    lines = []
    for paragraph in document.paragraphs:
        # Word keeps pictures inline, so their position is exact.
        for marker in _docx_paragraph_images(paragraph, document):
            lines.append(marker)

        text = paragraph.text.strip()
        if not text:
            continue
        # A fully bold option line is a common way of marking the answer;
        # normalise it to the asterisk convention the parser understands.
        runs = [run for run in paragraph.runs if run.text.strip()]
        if runs and all(run.bold for run in runs) and OPTION_RE.match(text):
            text = f"{text} *"
        lines.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    text = "\n".join(lines).strip()
    if not text:
        raise DocumentError("That Word document appears to be empty.")
    return text


def _docx_paragraph_images(paragraph, document):
    """Save any pictures anchored in this paragraph and return their markers."""
    markers = []
    namespace = (
        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
    )
    for blip in paragraph._element.iter(
        "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
    ):
        rel_id = blip.get(namespace)
        if not rel_id:
            continue
        try:
            part = document.part.related_parts[rel_id]
            data = part.blob
            if not is_meaningful_image(data):
                continue
            extension = (part.partname.ext or "png").lstrip(".")
            markers.append(f"[[IMAGE:{store_image(data, extension)}]]")
        except Exception:
            continue
    return markers


# ---------------------------------------------------------------------------
# Rule-based question reader
# ---------------------------------------------------------------------------


def _clean(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def _blank_question():
    return {
        "type": "single",
        "text": "",
        "hint": "",
        "points": 1,
        "image": "",
        "image_caption": "",
        "accepted_answers": "",
        "explanation": "",
        "choices": [],
        "_answer_line": "",
        "_number": None,
    }


def _letters(value: str):
    """Pull answer letters out of 'B', 'b and d', 'A, C' …"""
    head = re.split(r"\b(?:because|since|as it|explanation)\b", value, 1, re.I)[0]
    tokens = re.findall(r"\b([a-h])\b", head, re.I)
    return [t.upper() for t in tokens]


def _finalise(question, key_map):
    """Decide the question type and apply the answer key."""
    question["text"] = _clean(question["text"])
    if not question["text"]:
        return None

    marks = MARKS_RE.search(question["text"])
    if marks:
        question["points"] = max(1, int(marks.group(1)))
        question["text"] = _clean(MARKS_RE.sub("", question["text"]))

    answer_line = question.pop("_answer_line", "")
    number = question.pop("_number", None)
    if not answer_line and number and number in key_map:
        answer_line = key_map[number]

    choices = question["choices"]

    # No options: a typed short answer.
    if not choices:
        if not answer_line:
            return None
        question["type"] = "short_text"
        question["accepted_answers"] = _clean(answer_line)
        return question

    letters = _letters(answer_line) if answer_line else []
    if letters:
        for index, choice in enumerate(choices):
            if chr(65 + index) in letters:
                choice["is_correct"] = True

    # Fall back to matching the answer text against the option text.
    if not any(c["is_correct"] for c in choices) and answer_line:
        target = _clean(answer_line).lower()
        for choice in choices:
            if choice["text"].lower() == target:
                choice["is_correct"] = True
                break

    if not any(c["is_correct"] for c in choices):
        return None  # unusable without a key — reported as skipped

    labels = {c["text"].strip().lower() for c in choices}
    if len(choices) == 2 and labels <= {"true", "false"}:
        question["type"] = "true_false"
    elif sum(1 for c in choices if c["is_correct"]) > 1:
        question["type"] = "multiple"
    else:
        question["type"] = "single"

    return question


def parse_questions(text: str):
    """Read questions out of plain text. Returns (questions, stats)."""
    lines = [ln.strip() for ln in text.replace("\r", "").split("\n")]
    lines = [ln for ln in lines if ln and not PAGE_NOISE_RE.match(ln)]

    # An answer key printed at the end of the paper.
    key_map, body = {}, lines
    for index, line in enumerate(lines):
        if KEY_HEADER_RE.match(line):
            for entry in lines[index + 1 :]:
                for number, letters in KEY_ENTRY_RE.findall(entry):
                    key_map[int(number)] = letters
            if key_map:
                body = lines[:index]
            break

    questions, current, skipped = [], None, 0
    orphan_image = ""  # a picture printed before its question stem

    def flush():
        nonlocal current, skipped
        if not current:
            return
        finished = _finalise(current, key_map)
        if finished:
            questions.append(finished)
        else:
            skipped += 1
        current = None

    for line in body:
        image_match = IMAGE_MARKER_RE.match(line)
        if image_match:
            path = image_match.group(1)
            # Attach to the question being read; otherwise hold it for the next.
            if current and not current["image"]:
                current["image"] = path
            elif not current:
                orphan_image = path
            continue

        question_match = QUESTION_RE.match(line)
        option_match = OPTION_RE.match(line)
        answer_match = ANSWER_RE.match(line)
        explain_match = EXPLANATION_RE.match(line)

        # "A." only starts an option when a question is already open, and a
        # numbered line only starts a question when it is not an option body.
        if question_match and not (current and option_match and len(current["choices"])):
            flush()
            current = _blank_question()
            current["_number"] = int(question_match.group(1))
            current["text"] = question_match.group(2)
            if orphan_image:
                current["image"] = orphan_image
                orphan_image = ""
            continue

        if not current:
            continue

        if explain_match:
            current["explanation"] = _clean(explain_match.group(1))
            continue

        if answer_match:
            current["_answer_line"] = answer_match.group(1)
            continue

        if option_match:
            body_text = option_match.group(2)
            marked = bool(CORRECT_MARK_RE.search(body_text))
            current["choices"].append(
                {"text": _clean(CORRECT_MARK_RE.sub("", body_text)), "is_correct": marked}
            )
            continue

        # Continuation line: extends the last option, or the question stem.
        if current["choices"]:
            current["choices"][-1]["text"] = _clean(f"{current['choices'][-1]['text']} {line}")
        else:
            current["text"] = f"{current['text']} {line}"

    flush()

    stats = {
        "found": len(questions),
        "skipped": skipped,
        "total_marks": sum(q["points"] for q in questions),
        "with_images": sum(1 for q in questions if q.get("image")),
    }
    return questions, stats


# ---------------------------------------------------------------------------
# Optional local LLM (Ollama) — free and offline; never required
# ---------------------------------------------------------------------------

LLM_PROMPT = """You extract exam questions from documents.
Return ONLY a JSON array. Each element must be:
{"text": str, "type": "single"|"multiple"|"true_false"|"short_text",
 "points": int, "choices": [{"text": str, "is_correct": bool}],
 "accepted_answers": str, "explanation": str}
Rules: use "choices" only for choice questions; mark every correct option;
for short_text put the answer in accepted_answers and leave choices empty;
never invent questions that are not in the document.

DOCUMENT:
"""


def _ollama_config():
    config = getattr(settings, "AI_IMPORT", {})
    return (
        config.get("PROVIDER", "auto"),
        config.get("OLLAMA_URL", "http://127.0.0.1:11434").rstrip("/"),
        config.get("OLLAMA_MODEL", "llama3.2"),
        int(config.get("TIMEOUT", 120)),
    )


def parse_with_ollama(text: str):
    """Ask a locally running Ollama model to structure the document."""
    _, url, model, timeout = _ollama_config()
    # The model only sees prose; picture markers are handled by the reader.
    clean_text = "\n".join(
        line for line in text.split("\n") if not IMAGE_MARKER_RE.match(line.strip())
    )
    payload = json.dumps(
        {
            "model": model,
            "prompt": LLM_PROMPT + clean_text[:24000],
            "stream": False,
            "format": "json",
            "options": {"temperature": 0},
        }
    ).encode()

    request = urllib.request.Request(
        f"{url}/api/generate", data=payload, headers={"Content-Type": "application/json"}
    )
    with urllib.request.urlopen(request, timeout=timeout) as response:
        body = json.loads(response.read().decode())

    content = body.get("response", "")
    match = re.search(r"[\[\{].*[\]\}]", content, re.S)
    if not match:
        raise ValueError("The model did not return JSON.")
    data = json.loads(match.group(0))
    if isinstance(data, dict):
        data = data.get("questions", [])

    questions = []
    for item in data:
        text_value = _clean(item.get("text", ""))
        if not text_value:
            continue
        choices = [
            {"text": _clean(c.get("text", "")), "is_correct": bool(c.get("is_correct"))}
            for c in item.get("choices", [])
            if _clean(c.get("text", ""))
        ]
        qtype = item.get("type", "single")
        if qtype not in {"single", "multiple", "true_false", "short_text"}:
            qtype = "single"
        if qtype != "short_text" and not any(c["is_correct"] for c in choices):
            continue
        questions.append(
            {
                "type": qtype,
                "text": text_value,
                "hint": "",
                "points": max(1, int(item.get("points", 1) or 1)),
                "image": "",
                "image_caption": "",
                "accepted_answers": _clean(item.get("accepted_answers", "")),
                "explanation": _clean(item.get("explanation", "")),
                "choices": [] if qtype == "short_text" else choices,
            }
        )
    if not questions:
        raise ValueError("The model returned no usable questions.")
    return questions


def ollama_available() -> bool:
    _, url, _, _ = _ollama_config()
    try:
        with urllib.request.urlopen(f"{url}/api/tags", timeout=2):
            return True
    except (urllib.error.URLError, OSError, ValueError):
        return False


def build_questions(text: str):
    """Extract questions, preferring a local model when one is available."""
    provider, *_ = _ollama_config()

    if provider in ("ollama", "auto") and (provider == "ollama" or ollama_available()):
        try:
            questions = parse_with_ollama(text)
            return questions, {
                "found": len(questions),
                "skipped": 0,
                "total_marks": sum(q["points"] for q in questions),
                "with_images": 0,
            }, "local-ai"
        except Exception:
            pass  # fall through to the offline reader

    questions, stats = parse_questions(text)
    return questions, stats, "document-parser"
